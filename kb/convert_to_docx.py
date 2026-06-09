#!/usr/bin/env python3
"""
Markdown -> DOCX converter tuned for GoHighLevel Knowledge Base ingestion.

GHL's KB chunker splits documents on H1-H4 heading boundaries. This converter
preserves heading hierarchy exactly so each section becomes its own clean chunk.

Usage:
    python3 convert_to_docx.py            # converts every .md in this folder to upload/<name>.docx
    python3 convert_to_docx.py 14_FAQ_MASTER.md   # converts a single file

Output:
    kb/upload/*.docx   <-- ready to upload to GHL Knowledge Base
    kb/upload/*.csv    <-- CSVs copied as-is (GHL accepts them directly)
"""

import os
import re
import sys
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent
UPLOAD = HERE / "upload"


def parse_markdown_to_docx(md_text: str, out_path: Path) -> None:
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    in_table = False
    table_rows: list[list[str]] = []
    in_code = False

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph(line)
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            i += 1
            continue

        # Tables (markdown |col|col| style)
        if line.strip().startswith("|") and line.strip().endswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # skip separator row like |---|---|
            if all(re.fullmatch(r":?-+:?", c) for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                # flush table
                if table_rows:
                    n_cols = max(len(r) for r in table_rows)
                    tbl = doc.add_table(rows=len(table_rows), cols=n_cols)
                    tbl.style = 'Light Grid Accent 1'
                    for r, row in enumerate(table_rows):
                        for c in range(n_cols):
                            cell_text = row[c] if c < len(row) else ""
                            tbl.cell(r, c).text = strip_inline_md(cell_text)
                in_table = False
                table_rows = []

        # Headings
        if line.startswith("###### "):
            doc.add_heading(strip_inline_md(line[7:]), level=6)
        elif line.startswith("##### "):
            doc.add_heading(strip_inline_md(line[6:]), level=5)
        elif line.startswith("#### "):
            doc.add_heading(strip_inline_md(line[5:]), level=4)
        elif line.startswith("### "):
            doc.add_heading(strip_inline_md(line[4:]), level=3)
        elif line.startswith("## "):
            doc.add_heading(strip_inline_md(line[3:]), level=2)
        elif line.startswith("# "):
            doc.add_heading(strip_inline_md(line[2:]), level=1)
        # Horizontal rule
        elif line.strip() == "---":
            doc.add_paragraph()
        # Blockquote
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(strip_inline_md(line[2:]))
            run.italic = True
        # Unordered list
        elif re.match(r"^\s*[-*]\s+", line):
            indent = len(line) - len(line.lstrip())
            content = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(strip_inline_md(content), style='List Bullet')
            if indent >= 2:
                p.paragraph_format.left_indent = Pt(18 + (indent // 2) * 18)
        # Ordered list
        elif re.match(r"^\s*\d+\.\s+", line):
            content = re.sub(r"^\s*\d+\.\s+", "", line)
            doc.add_paragraph(strip_inline_md(content), style='List Number')
        # Blank line
        elif line.strip() == "":
            doc.add_paragraph()
        # Plain paragraph
        else:
            add_runs_with_inline_formatting(doc, line)

        i += 1

    # Flush any final table
    if in_table and table_rows:
        n_cols = max(len(r) for r in table_rows)
        tbl = doc.add_table(rows=len(table_rows), cols=n_cols)
        tbl.style = 'Light Grid Accent 1'
        for r, row in enumerate(table_rows):
            for c in range(n_cols):
                cell_text = row[c] if c < len(row) else ""
                tbl.cell(r, c).text = strip_inline_md(cell_text)

    doc.save(out_path)


INLINE_PATTERNS = [
    (re.compile(r"\*\*([^*]+)\*\*"), "bold"),
    (re.compile(r"__([^_]+)__"), "bold"),
    (re.compile(r"\*([^*]+)\*"), "italic"),
    (re.compile(r"_([^_]+)_"), "italic"),
    (re.compile(r"`([^`]+)`"), "code"),
]


def add_runs_with_inline_formatting(doc, line: str) -> None:
    p = doc.add_paragraph()
    # Strip links: [text](url) -> text (url)
    line = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", line)
    # Greedy: handle bold first, then italic, then code
    add_segment(p, line)


def add_segment(p, text: str) -> None:
    # Try bold **...**
    m = re.search(r"\*\*([^*]+)\*\*", text)
    if m:
        before, mid, after = text[:m.start()], m.group(1), text[m.end():]
        if before:
            add_segment(p, before)
        r = p.add_run(strip_inline_md(mid))
        r.bold = True
        if after:
            add_segment(p, after)
        return
    # Italic *...*
    m = re.search(r"(?<!\*)\*([^*]+)\*(?!\*)", text)
    if m:
        before, mid, after = text[:m.start()], m.group(1), text[m.end():]
        if before:
            add_segment(p, before)
        r = p.add_run(mid)
        r.italic = True
        if after:
            add_segment(p, after)
        return
    # Code `...`
    m = re.search(r"`([^`]+)`", text)
    if m:
        before, mid, after = text[:m.start()], m.group(1), text[m.end():]
        if before:
            add_segment(p, before)
        r = p.add_run(mid)
        r.font.name = 'Consolas'
        r.font.size = Pt(10)
        if after:
            add_segment(p, after)
        return
    # No formatting
    p.add_run(text)


def strip_inline_md(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text


def main() -> None:
    UPLOAD.mkdir(exist_ok=True)

    if len(sys.argv) > 1:
        files = [HERE / arg for arg in sys.argv[1:]]
    else:
        files = sorted(HERE.glob("*.md")) + sorted(HERE.glob("*.csv"))

    for src in files:
        if src.suffix == ".md":
            # skip the owner README from the upload pile
            if src.name == "00_README_FOR_OWNER.md":
                continue
            # skip the system prompt — that goes into the agent config, not the KB
            if src.name == "01_AGENT_SYSTEM_PROMPT.md":
                continue
            out = UPLOAD / f"{src.stem}.docx"
            md_text = src.read_text(encoding="utf-8")
            parse_markdown_to_docx(md_text, out)
            print(f"  -> {out.name}")
        elif src.suffix == ".csv":
            out = UPLOAD / src.name
            shutil.copy(src, out)
            print(f"  -> {out.name} (copied)")

    print(f"\nDone. {len(list(UPLOAD.iterdir()))} files in {UPLOAD}")


if __name__ == "__main__":
    main()
